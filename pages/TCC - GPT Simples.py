__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')    

import os
import re
import fitz  # PyMuPDF para leitura de PDFs
from docx import Document as DocxDocument
import streamlit as st
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_openai import ChatOpenAI
import chromadb
from chromadb.config import Settings
import traceback
import subprocess

# Função para remover formatação Markdown do texto
def limpar_formatacao_markdown(texto):
    texto = re.sub(r'\*\*([^*]+)\*\*', r'\1', texto)
    texto = re.sub(r'\*([^*]+)\*', r'\1', texto)
    texto = re.sub(r'__([^_]+)__', r'\1', texto)
    texto = re.sub(r'_([^_]+)_', r'\1', texto)
    texto = re.sub(r'`([^`]+)`', r'\1', texto)
    return texto

# Função para salvar documento em formato .docx no PC do usuário
def salvar_documento_docx(tipo_documento, conteudo):
    user_home = os.path.expanduser("~")
    base_dir = os.path.join(user_home, "Downloads", "Artefatos")
    base_filename = f"{tipo_documento}.docx"
    caminho_docx = os.path.join(base_dir, base_filename)
    
    if os.path.exists(caminho_docx):
        i = 1
        while os.path.exists(caminho_docx):
            base_filename = f"{tipo_documento} ({i}).docx"
            caminho_docx = os.path.join(base_dir, base_filename)
            i += 1
    
    os.makedirs(os.path.dirname(caminho_docx), exist_ok=True)
    doc = DocxDocument()

    doc.add_heading(tipo_documento, level=1)

    for campo, resposta in conteudo.items():
        resposta_limpa = limpar_formatacao_markdown(resposta)
        doc.add_heading(campo, level=2)
        doc.add_paragraph(resposta_limpa, style='Normal')

    doc.save(caminho_docx)
    return caminho_docx

# Função para carregar arquivos PDF, DOCX e TXT
def carregar_arquivo(file_path):
    documents = []
    filename = os.path.basename(file_path)

    if filename.endswith(".pdf"):
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        documents.append(Document(page_content=text, metadata={"source": filename}))
    elif filename.endswith(".docx"):
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        documents.append(Document(page_content=text, metadata={"source": filename}))
    elif filename.endswith(".txt"):
        codificacoes = ['utf-8', 'latin-1', 'cp1252']
        for encoding in codificacoes:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                documents.append(Document(page_content=text, metadata={"source": filename}))
                break
            except UnicodeDecodeError:
                continue
        else:
            raise RuntimeError(f"Não foi possível carregar o arquivo: {file_path} com as codificações {codificacoes}")
    else:
        raise ValueError(f"Formato de arquivo não suportado: {file_path}")

    return documents

def reinicializar_chain():
    return ConversationalRetrievalChain.from_llm(
        llm=chat_model,
        chain_type="map_reduce",
        retriever=db.as_retriever(return_source_documents=True, top_k=5),
        memory=ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    )

def anonimizar_nomes(texto):
    nomes_regex = r'\b[A-Z][a-z]+(?:\s[A-Z][a-z]+)*\b'
    texto_anonimizado = re.sub(nomes_regex, '[NOME]', texto)
    return texto_anonimizado

def anonimizar_emails(texto):
    email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    texto_anonimizado = re.sub(email_regex, '[EMAIL]', texto)
    return texto_anonimizado

def anonimizar_enderecos(texto):
    enderecos_regex = r'\d{1,4} [A-Z][a-z]+(?: [A-Z][a-z]+)*(?:, [A-Z]{2})? \d{5}'
    texto_anonimizado = re.sub(enderecos_regex, '[ENDEREÇO]', texto)
    return texto_anonimizado

def anonimizar_texto(texto):
    texto = anonimizar_nomes(texto)
    texto = anonimizar_emails(texto)
    texto = anonimizar_enderecos(texto)
    return texto

templates = {
    "ETP": {
        "1. DESCRIÇÃO DA NECESSIDADE DA CONTRATAÇÃO": "Descreva de forma detalhada a necessidade que leva à contratação, incluindo a situação atual, os problemas enfrentados e o impacto no bem-estar e interesse público. Quantifique e qualifique o problema, mencionando as unidades envolvidas e os esforços anteriores para resolver a questão.",
        "2. PREVISÃO DA CONTRATAÇÃO NO PLANO DE CONTRATAÇÕES ANUAL – PCA": "Indique se a contratação está prevista no Plano de Contratações Anual (PCA) do MPBA. Caso não esteja, justifique a ausência, fornecendo contexto e alternativas planejadas.",
    },
    "TR": {
        "1. OBJETO": "Descreva detalhadamente o objeto da contratação, incluindo condições, quantidades e especificações técnicas estabelecidas neste Termo de Referência e seus anexos.",
    }
}

def preencher_documento(tipo_documento, retrieval_chain_config):
    inicial_instrução = """
    Considere que todo conteúdo gerado deve atender aos requisitos do Ministério Público do Estado da Bahia,
    conforme as diretrizes e regulamentações estabelecidas pela Lei 14.133/2021. Ao preencher cada campo,
    leve em conta as particularidades desta legislação e a importância de alinhamento com o Plano de Contratações Anual (PCA)
    e outros normativos relevantes para licitações e contratos administrativos.
    As referências devem ser contextualizadas para o MPBA e suas necessidades específicas.
    """
    if tipo_documento not in templates:
        raise ValueError(f"Tipo de documento {tipo_documento} não é suportado.")

    template = templates[tipo_documento]

    for campo, descricao in template.items():
        question = inicial_instrução + f" Preencha o campo '{campo}' de acordo com a seguinte descrição orientativa: {descricao}. Certifique-se de que o preenchimento esteja em conformidade com a Lei 14.133/2021 e as diretrizes do Ministério Público do Estado da Bahia."
        response = retrieval_chain_config.invoke({"question": question})
        template[campo] = response['answer']

    return template

def preencher_sequencia_documentos(retrieval_chain_config, tipo_documento_selecionado):
    documento_preenchido = preencher_documento(tipo_documento_selecionado, retrieval_chain_config)
    caminho_salvo = salvar_documento(tipo_documento_selecionado, documento_preenchido)
    return caminho_salvo

def salvar_documento(tipo_documento, conteudo):
    conteudo_anonimizado = {campo: anonimizar_texto(resposta) for campo, resposta in conteudo.items()}
    return salvar_documento_docx(tipo_documento, conteudo)

def iniciar_processo(uploaded_files):
    global retrieval_chain_config, chat_model, db

    try:
        documentos = []
        for uploaded_file in uploaded_files:
            with open(uploaded_file.name, "wb") as f:
                f.write(uploaded_file.getbuffer())
            documentos.extend(carregar_arquivo(uploaded_file.name))

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=50)
        docs = text_splitter.split_documents(documentos)

        os.environ["OPENAI_API_KEY"] = st.secrets["KEY"]

        embedder = OpenAIEmbeddings(model="text-embedding-3-large")

        # Configuração manual do cliente Chroma
        chroma_settings = Settings(anonymized_telemetry=False)
        db = Chroma.from_documents(docs, embedder, client_settings=chroma_settings)

        chat_model = ChatOpenAI(temperature=0.5, model_name="gpt-4o")
        memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
        retrieval_chain_config = reinicializar_chain()
        
        st.success("Documentos carregados e processados com sucesso.")
        return True

    except Exception as e:
        st.error(f"Erro ao iniciar o processo: {str(e)}")
        st.error(traceback.format_exc())
        return False

def gerar_documento(retrieval_chain_config, tipo_documento_selecionado):
    try:
        caminho_salvo = preencher_sequencia_documentos(retrieval_chain_config, tipo_documento_selecionado)
        st.success("Documento gerado com sucesso.")
        st.session_state['caminho_salvo'] = caminho_salvo
        return caminho_salvo
    except Exception as e:
        st.error(f"Erro ao gerar documento: {str(e)}")
        st.error(traceback.format_exc())
        return None

def reset_app():
    st.experimental_set_query_params()  # Remove all query parameters, resetting the app
    st.experimental_rerun()

def reinstall_dependencies():
    try:
        # Lista de pacotes necessários
        packages = [
            "langchain",
            "openai",
            "python-docx",
            "docx2txt",
            "pysqlite3-binary",
            "langchain-community",
            "tiktoken",
            "langchain_openai",
            "langchain-chroma",
            "PyMuPDF",
            "chromadb",
            "streamlit",
            "tiktoken"
        ]

        # Instalar cada pacote
        for package in packages:
            result = subprocess.run([sys.executable, "-m", "pip", "install", "--upgrade", package], capture_output=True, text=True)
            if result.returncode != 0:
                st.error(f"Erro ao reinstalar {package}: {result.stderr}")
                return

        st.success("Dependências reinstaladas com sucesso. Reiniciando aplicação...")
        st.experimental_rerun()
    except Exception as e:
        st.error(f"Erro ao reinstalar dependências: {str(e)}")
        st.error(traceback.format_exc())

st.title("Gerador de Artefatos de Licitação do MPBA")

# Botão para resetar a aplicação
if st.button("Resetar Aplicação"):
    reset_app()

# Botão para reinstalar dependências
if st.button("Reinstalar Dependências"):
    reinstall_dependencies()

# Upload de arquivos
uploaded_files = st.file_uploader("Carregue seus arquivos", accept_multiple_files=True, type=["pdf", "docx", "txt"])

if uploaded_files:
    if iniciar_processo(uploaded_files):
        st.write("Arquivos carregados:")
        for file in uploaded_files:
            st.write(file.name)

        # Dropdown para selecionar o tipo de documento
        tipo_documento_selecionado = st.selectbox("Selecione o tipo de documento", list(templates.keys()))

        if st.button("Gerar Documento"):
            with st.spinner('Gerando documento, por favor aguarde...'):
                caminho_salvo = gerar_documento(retrieval_chain_config, tipo_documento_selecionado)

            if caminho_salvo:
                with open(caminho_salvo, "rb") as file:
                    btn = st.download_button(
                        label="Download Documento",
                        data=file,
                        file_name=os.path.basename(caminho_salvo),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
